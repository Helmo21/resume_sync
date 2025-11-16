"""
Document Generation Service for Resumes

Generates PDF and DOCX files from CV data using ATS-friendly templates.
Supports both custom generation and using pre-made DOCX templates.
"""

from typing import Dict, List, Optional
from datetime import datetime
import os
import re
from pathlib import Path

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfgen import canvas
from reportlab.graphics.shapes import Drawing, Circle
from reportlab.graphics import renderPDF
import io
from PIL import Image as PILImage
import requests

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ATSTemplateGenerator:
    """
    Generates ATS-friendly resumes in PDF and DOCX formats.
    ATS (Applicant Tracking Systems) require simple, clean formatting.

    Supports both custom generation and using pre-made DOCX templates.
    """

    def __init__(self, templates_dir: Optional[str] = None):
        """
        Initialize template generator.

        Args:
            templates_dir: Path to directory containing DOCX templates
        """
        self.templates = {
            "modern": self._modern_template,
            "classic": self._classic_template,
            "technical": self._technical_template
        }

        # Default templates directory
        self.templates_dir = templates_dir or "/home/antoine/Documents/dev/ResumeSync/teamplate"

        # Available DOCX templates
        self.docx_templates = {
            "accounting": "ATS Bold accounting resume.docx",
            "office_manager": "ATS office manager resume.docx",
            "attorney": "Attorney resume.docx",
            "sales": "Modern bold sales resume.docx"
        }

    def generate_pdf(
        self,
        resume_data: Dict,
        output_path: str,
        template: str = "modern"
    ) -> str:
        """
        Generate PDF resume.

        Args:
            resume_data: Structured resume data from CV generator
            output_path: Path to save PDF file
            template: Template name (modern, classic, technical)

        Returns:
            str: Path to generated PDF
        """
        template_func = self.templates.get(template, self._modern_template)
        return template_func(resume_data, output_path, format="pdf")

    def generate_docx(
        self,
        resume_data: Dict,
        output_path: str,
        template: str = "modern",
        use_template_file: bool = False
    ) -> str:
        """
        Generate DOCX resume.

        Args:
            resume_data: Structured resume data from CV generator
            output_path: Path to save DOCX file
            template: Template name (modern, classic, technical) or template key (accounting, sales, etc.)
            use_template_file: If True, use pre-made DOCX template file

        Returns:
            str: Path to generated DOCX
        """
        if use_template_file:
            # Use pre-made DOCX template
            return self._generate_from_docx_template(resume_data, output_path, template)
        else:
            # Generate from scratch
            template_func = self.templates.get(template, self._modern_template)
            return template_func(resume_data, output_path, format="docx")

    def _modern_template(self, resume_data: Dict, output_path: str, format: str = "pdf") -> str:
        """Modern ATS-friendly template"""
        if format == "pdf":
            return self._generate_modern_pdf(resume_data, output_path)
        else:
            return self._generate_modern_docx(resume_data, output_path)

    def _classic_template(self, resume_data: Dict, output_path: str, format: str = "pdf") -> str:
        """Classic ATS-friendly template"""
        if format == "pdf":
            return self._generate_classic_pdf(resume_data, output_path)
        else:
            return self._generate_classic_docx(resume_data, output_path)

    def _technical_template(self, resume_data: Dict, output_path: str, format: str = "pdf") -> str:
        """Technical ATS-friendly template"""
        if format == "pdf":
            return self._generate_technical_pdf(resume_data, output_path)
        else:
            return self._generate_technical_docx(resume_data, output_path)

    def _create_circular_image(self, image_url: str, size: float = 1.0*inch) -> Optional[Image]:
        """
        Download and create a circular profile image.

        Args:
            image_url: URL or path to the image
            size: Diameter of the circular image in inches

        Returns:
            ReportLab Image object or None if failed
        """
        try:
            # Download or load image
            if image_url.startswith('http'):
                response = requests.get(image_url, timeout=10)
                image_data = io.BytesIO(response.content)
                img = PILImage.open(image_data)
            else:
                img = PILImage.open(image_url)

            # Convert to RGB if necessary
            if img.mode != 'RGB':
                img = img.convert('RGB')

            # Create circular mask
            width, height = img.size
            min_dim = min(width, height)

            # Crop to square
            left = (width - min_dim) // 2
            top = (height - min_dim) // 2
            right = left + min_dim
            bottom = top + min_dim
            img = img.crop((left, top, right, bottom))

            # Create circular mask
            mask = PILImage.new('L', (min_dim, min_dim), 0)
            from PIL import ImageDraw
            draw = ImageDraw.Draw(mask)
            draw.ellipse((0, 0, min_dim, min_dim), fill=255)

            # Apply mask
            output = PILImage.new('RGBA', (min_dim, min_dim), (255, 255, 255, 0))
            output.paste(img, (0, 0))
            output.putalpha(mask)

            # Resize to desired size
            pixel_size = int(size * 72)  # Convert inches to pixels (72 DPI)
            output = output.resize((pixel_size, pixel_size), PILImage.Resampling.LANCZOS)

            # Save to BytesIO
            img_buffer = io.BytesIO()
            output.save(img_buffer, format='PNG')
            img_buffer.seek(0)

            # Create ReportLab Image
            return Image(img_buffer, width=size, height=size)

        except Exception as e:
            print(f"Error creating circular image: {e}")
            return None

    def _generate_modern_pdf(self, resume_data: Dict, output_path: str) -> str:
        """Generate modern professional PDF matching high-quality reference design (1 page optimized)"""
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                                rightMargin=0.4*inch, leftMargin=0.4*inch,
                                topMargin=0.4*inch, bottomMargin=0.4*inch)

        story = []
        styles = getSampleStyleSheet()

        # Custom styles - Compact & Professional (optimized for 1-page)
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#000000'),
            spaceAfter=2,
            alignment=TA_LEFT,
            fontName='Helvetica-Bold',
            leading=23
        )

        subtitle_style = ParagraphStyle(
            'SubtitleStyle',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#555555'),
            spaceAfter=4,
            alignment=TA_LEFT,
            fontName='Helvetica-Oblique',
            leading=11
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            spaceAfter=3,
            spaceBefore=6,
            fontName='Helvetica-Bold',
            leading=11,
            borderWidth=0,
            borderPadding=0,
            leftIndent=0
        )

        # Header styles for contact info (optimized for 1-page)
        contact_style = ParagraphStyle(
            'ContactStyle',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#555555'),
            spaceAfter=2,
            alignment=TA_RIGHT,
            leading=10
        )

        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=8.5,
            textColor=colors.HexColor('#333333'),
            spaceAfter=2,
            alignment=TA_JUSTIFY,
            leading=10
        )

        job_title_style = ParagraphStyle(
            'JobTitleStyle',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            spaceAfter=1,
            fontName='Helvetica-Bold',
            leading=10
        )

        job_meta_style = ParagraphStyle(
            'JobMetaStyle',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#666666'),
            spaceAfter=2,
            fontName='Helvetica-Oblique',
            leading=9
        )

        # Header avec nom, titre et contact infos alignés
        personal_info = resume_data.get('personal_info', {})
        full_name = personal_info.get('full_name', 'N/A').upper()

        # Try to get profile photo
        profile_image = None
        photo_url = personal_info.get('photo_url') or personal_info.get('profile_image_url')
        if photo_url:
            profile_image = self._create_circular_image(photo_url, size=1.0*inch)

        # Colonne gauche : Nom + Headline (job title)
        left_column = []
        left_column.append(Paragraph(full_name, title_style))

        # Job title variants if available - below the name (for ATS optimization)
        headline_text = ""
        job_title_variants = personal_info.get('job_title_variants', [])
        if job_title_variants:
            # Use job title variants separated by pipes
            headline_text = " | ".join(job_title_variants)
        elif resume_data.get('headline'):
            headline_text = resume_data['headline']
        elif personal_info.get('headline'):
            headline_text = personal_info['headline']

        if headline_text:
            left_column.append(Paragraph(headline_text.upper(), subtitle_style))

        # Colonne droite : Photo + Contact info
        contact_info_right = []

        # Add profile photo if available
        if profile_image:
            contact_info_right.append(profile_image)

        # Add phone and email below photo (or standalone if no photo)
        contact_text_items = []
        if personal_info.get('phone'):
            contact_text_items.append(Paragraph(personal_info['phone'], contact_style))
        if personal_info.get('email'):
            contact_text_items.append(Paragraph(personal_info['email'], contact_style))

        # Add additional links (GitHub, portfolio, etc.)
        for link in personal_info.get('additional_links', []):
            if link:
                contact_text_items.append(Paragraph(link, contact_style))

        # Create header table with 3 columns: name/title (left), spacer (middle), photo/contact (right)
        if profile_image or contact_text_items:
            # Combine photo and contact text in right column
            right_column_content = contact_info_right + contact_text_items

            header_data = [[left_column, right_column_content]]
            header_table = Table(header_data, colWidths=[4.5*inch, 2.5*inch])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ALIGN', (0, 0), (0, 0), 'LEFT'),
                ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ]))
            story.append(header_table)
        else:
            # Sinon juste le nom et titre
            for item in left_column:
                story.append(item)

        # Horizontal separator line after header (reduced spacing)
        story.append(Spacer(1, 0.03*inch))
        separator = Table([['']], colWidths=[7.5*inch])
        separator.setStyle(TableStyle([
            ('LINEABOVE', (0, 0), (-1, -1), 1.5, colors.HexColor('#000000')),
        ]))
        story.append(separator)
        story.append(Spacer(1, 0.05*inch))

        # Professional Summary
        if resume_data.get('professional_summary'):
            story.append(Paragraph("<i>PROFESSIONAL SUMMARY</i>", heading_style))
            story.append(Spacer(1, 0.02*inch))

            summary_text = resume_data['professional_summary']
            story.append(Paragraph(summary_text, body_style))
            story.append(Spacer(1, 0.05*inch))

        # Work Experience (reduced spacing for 1-page)
        if resume_data.get('work_experience'):
            story.append(Paragraph("<i>WORK HISTORY</i>", heading_style))
            story.append(Spacer(1, 0.02*inch))

            for i, exp in enumerate(resume_data['work_experience']):
                # Job title and dates on same line
                title = exp.get('title', 'N/A')
                start_date = exp.get('start_date', '')
                end_date = exp.get('end_date', 'Present')
                date_range = f"{start_date} - {end_date}" if start_date else end_date

                # Title line with job title and dates
                title_line = f"<b>{title}, {date_range}</b>"
                story.append(Paragraph(title_line, job_title_style))

                # Company and location line
                company_parts = []
                if exp.get('company'):
                    company_parts.append(exp['company'])
                if exp.get('location'):
                    company_parts.append(exp['location'])

                if company_parts:
                    company_line = " – ".join(company_parts)
                    story.append(Paragraph(f"<i>{company_line}</i>", job_meta_style))

                # Bullets (NEW format from ContentWriter) - priority
                if exp.get('bullets'):
                    story.append(Spacer(1, 0.01*inch))
                    for bullet in exp['bullets']:
                        bullet_text = f"• {bullet}"
                        story.append(Paragraph(bullet_text, body_style))

                # Description as narrative paragraph (fallback for old format)
                elif exp.get('description'):
                    description = exp['description']
                    story.append(Spacer(1, 0.01*inch))
                    story.append(Paragraph(description, body_style))

                # Achievements as bullet points (fallback)
                elif exp.get('achievements'):
                    story.append(Spacer(1, 0.01*inch))
                    for achievement in exp['achievements']:
                        bullet_text = f"• {achievement}"
                        story.append(Paragraph(bullet_text, body_style))

                story.append(Spacer(1, 0.04*inch))

        # Skills Section (reduced spacing for 1-page)
        if resume_data.get('skills'):
            story.append(Paragraph("<i>SKILLS</i>", heading_style))
            story.append(Spacer(1, 0.02*inch))

            skills_dict = resume_data['skills']

            # Organize skills into categories with bullet points
            skill_items = []
            if isinstance(skills_dict, dict):
                # Skills organized by category - create bullet points with bold category names
                for category, skills_data in skills_dict.items():
                    if not skills_data or category.lower() == 'other':
                        continue

                    # Format category title nicely
                    category_title = category.replace('_', ' ').title()

                    # Handle different skill data formats
                    if isinstance(skills_data, dict):
                        # Nested dictionary: category -> {skill_name: detailed_skills}
                        # Use the detailed VALUES, not the generic keys
                        skills_list = list(skills_data.values())
                        skills_text = ", ".join(skills_list)
                    elif isinstance(skills_data, list):
                        # List of skills
                        skills_text = ", ".join(skills_data)
                    else:
                        # String or other format
                        skills_text = str(skills_data)

                    # Create bullet point with bold category name
                    skill_item = f"• <b>{category_title}:</b> {skills_text}"
                    skill_items.append(Paragraph(skill_item, body_style))
            else:
                # Skills as flat list
                skills_list = skills_dict if isinstance(skills_dict, list) else []
                for skill in skills_list:
                    skill_items.append(Paragraph(f"• {skill}", body_style))

            # Display skills in single column with minimal spacing
            for item in skill_items:
                story.append(item)
                # Minimal spacing after each skill category
                story.append(Spacer(1, 0.02*inch))

            story.append(Spacer(1, 0.05*inch))

        # Education (Formation) - reduced spacing for 1-page
        if resume_data.get('education'):
            story.append(Paragraph("<i>FORMATION</i>", heading_style))
            story.append(Spacer(1, 0.02*inch))

            for edu in resume_data['education']:
                degree = edu.get('degree', 'N/A')
                location = edu.get('location', '')

                # Format dates
                start_date = edu.get('start_date', '')
                end_date = edu.get('end_date', '')
                grad_year = edu.get('graduation_year', '')
                grad_date = edu.get('graduation_date', '')

                # Build date range
                date_range = ""
                if start_date and end_date:
                    date_range = f"{start_date}-{end_date}"
                elif grad_date:
                    date_range = grad_date
                elif grad_year:
                    date_range = grad_year

                # Degree title (bold) - on one line
                story.append(Paragraph(f"<b>{degree}</b>", job_title_style))

                # Location and date range on separate lines
                if location:
                    story.append(Paragraph(location, job_meta_style))

                if date_range:
                    story.append(Paragraph(f"<i>{date_range}</i>", job_meta_style))

                story.append(Spacer(1, 0.03*inch))

        # Certifications Section (reduced spacing for 1-page)
        if resume_data.get('certifications'):
            story.append(Paragraph("<i>CERTIFICATIONS</i>", heading_style))
            story.append(Spacer(1, 0.02*inch))

            for cert in resume_data['certifications']:
                if isinstance(cert, dict):
                    cert_name = cert.get('name', 'N/A')
                    cert_issuer = cert.get('issuer', '')
                    cert_date = cert.get('date', '')
                    cert_status = cert.get('status', '')

                    # Format: "Certification Name | Issuer | Date | Status"
                    cert_parts = [cert_name]
                    if cert_issuer:
                        cert_parts.append(cert_issuer)
                    if cert_date:
                        cert_parts.append(cert_date)
                    if cert_status and cert_status != 'Completed':
                        cert_parts.append(cert_status)

                    cert_text = " | ".join(cert_parts)
                    story.append(Paragraph(f"• {cert_text}", body_style))
                else:
                    story.append(Paragraph(f"• {str(cert)}", body_style))

            story.append(Spacer(1, 0.04*inch))

        # Projects Section (reduced spacing for 1-page)
        if resume_data.get('projects'):
            story.append(Paragraph("<i>PROJECTS</i>", heading_style))
            story.append(Spacer(1, 0.02*inch))

            for project in resume_data['projects']:
                if isinstance(project, dict):
                    project_name = project.get('name', 'N/A')
                    technologies = project.get('technologies', [])
                    description = project.get('description', '')
                    github_url = project.get('github_url', '')

                    # Project name (bold)
                    story.append(Paragraph(f"<b>{project_name}</b>", job_title_style))

                    # Technologies
                    if technologies:
                        tech_text = ", ".join(technologies) if isinstance(technologies, list) else str(technologies)
                        story.append(Paragraph(f"<i>Technologies: {tech_text}</i>", job_meta_style))

                    # Description
                    if description:
                        story.append(Paragraph(description, body_style))

                    # GitHub URL if available
                    if github_url:
                        story.append(Paragraph(f"<i>{github_url}</i>", job_meta_style))

                    story.append(Spacer(1, 0.03*inch))

        # Languages Section (reduced spacing for 1-page)
        if resume_data.get('languages'):
            story.append(Paragraph("<i>LANGUAGES</i>", heading_style))
            story.append(Spacer(1, 0.02*inch))

            language_items = []
            for lang in resume_data['languages']:
                if isinstance(lang, dict):
                    language = lang.get('language', 'N/A')
                    proficiency = lang.get('proficiency', '')
                    lang_text = f"{language}: {proficiency}" if proficiency else language
                    language_items.append(lang_text)
                else:
                    language_items.append(str(lang))

            # Display languages on one line, comma-separated
            if language_items:
                languages_text = " | ".join(language_items)
                story.append(Paragraph(languages_text, body_style))

            story.append(Spacer(1, 0.06*inch))

        # Build PDF
        doc.build(story)
        return output_path

    def _generate_modern_docx(self, resume_data: Dict, output_path: str) -> str:
        """
        Generate ATS-OPTIMIZED DOCX resume following ALL best practices:
        - No tables, images, text boxes, or complex formatting
        - Standard Word styles only (Heading 1, Heading 2, Normal)
        - Simple bullets, left-aligned text
        - Consistent date formatting
        - Standard fonts (Calibri 11pt)
        - Proper metadata
        - Keywords from job posting
        - Quantified results
        """
        doc = Document()

        # ===== DOCUMENT SETUP =====
        # Set document margins (1 inch = 2.54cm standard)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Set default font for entire document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Configure Heading 1 style (for section headers)
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = 'Calibri'
        heading1_style.font.size = Pt(12)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # Configure Heading 2 style (for job titles, degrees)
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Calibri'
        heading2_style.font.size = Pt(11)
        heading2_style.font.bold = True
        heading2_style.font.color.rgb = RGBColor(0, 0, 0)  # Black

        # ===== HEADER: NAME AND CONTACT =====
        personal_info = resume_data.get('personal_info', {})

        # Full name (larger, bold, left-aligned - NOT centered)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal_info.get('full_name', '').upper())
        name_run.font.name = 'Calibri'
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Job title variants (NEW - for ATS optimization)
        job_title_variants = personal_info.get('job_title_variants', [])
        if job_title_variants:
            job_title_para = doc.add_paragraph()
            job_title_run = job_title_para.add_run(" | ".join(job_title_variants))
            job_title_run.font.name = 'Calibri'
            job_title_run.font.size = Pt(12)
            job_title_run.font.bold = False
            job_title_run.font.color.rgb = RGBColor(80, 80, 80)
            job_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Contact info line (phone | email | city | LinkedIn | additional links)
        contact_parts = []
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('location'):
            # City only, not full address
            city = personal_info['location'].split(',')[0].strip()
            contact_parts.append(city)
        if personal_info.get('linkedin'):
            # Clean LinkedIn URL
            linkedin_url = personal_info['linkedin']
            if 'linkedin.com/in/' in linkedin_url:
                linkedin_url = linkedin_url.split('linkedin.com/in/')[-1].split('/')[0]
                contact_parts.append(f"linkedin.com/in/{linkedin_url}")
            else:
                contact_parts.append(linkedin_url)

        # Add additional links (GitHub, portfolio, etc.)
        for link in personal_info.get('additional_links', []):
            if link:
                # Clean up URL for display (remove https://)
                display_link = link.replace('https://', '').replace('http://', '')
                contact_parts.append(display_link)

        contact_para = doc.add_paragraph(' | '.join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add spacing after header
        doc.add_paragraph()

        # ===== PROFESSIONAL SUMMARY =====
        if resume_data.get('professional_summary'):
            doc.add_heading('PROFESSIONAL SUMMARY', level=1)
            summary_para = doc.add_paragraph(resume_data['professional_summary'])
            summary_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph()  # Spacing

        # ===== WORK EXPERIENCE =====
        if resume_data.get('work_experience'):
            doc.add_heading('WORK EXPERIENCE', level=1)

            for exp in resume_data['work_experience']:
                # Job Title (Bold, Heading 2 style)
                title = exp.get('title', 'N/A')
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(11)
                title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Company Name | City, Country (or State)
                company_parts = []
                if exp.get('company'):
                    company_parts.append(exp['company'])
                if exp.get('location'):
                    company_parts.append(exp['location'])

                if company_parts:
                    company_para = doc.add_paragraph(' | '.join(company_parts))
                    company_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Dates (consistent format: Mon YYYY - Mon YYYY or Present)
                start_date = exp.get('start_date', '')
                end_date = exp.get('end_date', 'Present')

                # Clean date format (remove duration like "· 2 yrs 1 mo")
                if ' · ' in end_date:
                    end_date = end_date.split(' · ')[0].strip()

                date_range = f"{start_date} - {end_date}" if start_date else end_date
                date_para = doc.add_paragraph(date_range)
                date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Bullets (NEW format from ContentWriter) - priority
                if exp.get('bullets'):
                    for bullet in exp['bullets']:
                        bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                        bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Description or Achievements (fallback for old format)
                elif exp.get('description'):
                    # Split description into sentences and create bullets
                    description = exp['description']
                    # Split by period, create max 5 bullet points
                    sentences = [s.strip() + '.' for s in description.split('.') if s.strip()]
                    for sentence in sentences[:5]:  # Max 5 bullets
                        if len(sentence) > 10:  # Skip very short fragments
                            bullet_para = doc.add_paragraph(sentence, style='List Bullet')
                            bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                elif exp.get('achievements'):
                    for achievement in exp.get('achievements', []):
                        bullet_para = doc.add_paragraph(achievement, style='List Bullet')
                        bullet_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Spacing between jobs
                doc.add_paragraph()

        # ===== EDUCATION =====
        if resume_data.get('education'):
            doc.add_heading('EDUCATION', level=1)

            for edu in resume_data['education']:
                # Degree - Specialization (Bold)
                degree_parts = []
                if edu.get('degree'):
                    degree_parts.append(edu['degree'])
                if edu.get('field') and edu.get('field') not in edu.get('degree', ''):
                    degree_parts.append(edu['field'])

                degree_text = ' - '.join(degree_parts) if degree_parts else 'N/A'
                degree_para = doc.add_paragraph()
                degree_run = degree_para.add_run(degree_text)
                degree_run.bold = True
                degree_run.font.size = Pt(11)
                degree_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # School Name | City, Country
                school_parts = []
                if edu.get('school'):
                    school_parts.append(edu['school'])
                if edu.get('location'):
                    school_parts.append(edu['location'])

                if school_parts:
                    school_para = doc.add_paragraph(' | '.join(school_parts))
                    school_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Graduation Date or Date Range
                date_parts = []
                if edu.get('start_date') and edu.get('end_date'):
                    date_parts.append(f"{edu['start_date']} - {edu['end_date']}")
                elif edu.get('graduation_date'):
                    date_parts.append(edu['graduation_date'])
                elif edu.get('graduation_year'):
                    date_parts.append(str(edu['graduation_year']))
                elif edu.get('end_date'):
                    date_parts.append(str(edu['end_date']))

                if date_parts:
                    date_para = doc.add_paragraph(date_parts[0])
                    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Add distinction if available
                if edu.get('distinction') or edu.get('honors'):
                    distinction = edu.get('distinction') or edu.get('honors')
                    dist_para = doc.add_paragraph(f"Distinction: {distinction}")
                    dist_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Spacing between degrees
                doc.add_paragraph()

        # ===== SKILLS =====
        if resume_data.get('skills'):
            doc.add_heading('SKILLS', level=1)

            skills_dict = resume_data['skills']

            # Handle different skill formats
            if isinstance(skills_dict, dict):
                for category, skills_data in skills_dict.items():
                    if not skills_data or category.lower() == 'other':
                        continue

                    # Format category nicely
                    category_title = category.replace('_', ' ').title()

                    # Extract detailed skill values
                    if isinstance(skills_data, dict):
                        # Nested dict: extract VALUES (detailed technical skills), not keys
                        skills_list = list(skills_data.values())
                    elif isinstance(skills_data, list):
                        skills_list = skills_data
                    else:
                        skills_list = [str(skills_data)]

                    # Create paragraph: "Category: Detailed technical skills"
                    if skills_list:
                        skills_para = doc.add_paragraph()
                        category_run = skills_para.add_run(f"{category_title}: ")
                        category_run.bold = True
                        skills_run = skills_para.add_run(", ".join(skills_list))
                        skills_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # Flat list of skills
                skills_list = skills_dict if isinstance(skills_dict, list) else []
                if skills_list:
                    skills_para = doc.add_paragraph(", ".join(skills_list))
                    skills_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ===== CERTIFICATIONS (Enhanced) =====
        if resume_data.get('certifications'):
            doc.add_paragraph()
            doc.add_heading('CERTIFICATIONS', level=1)

            for cert in resume_data['certifications']:
                if isinstance(cert, dict):
                    cert_name = cert.get('name', '')
                    cert_issuer = cert.get('issuer', cert.get('authority', ''))
                    cert_date = cert.get('date', '')
                    cert_status = cert.get('status', '')

                    cert_parts = [cert_name]
                    if cert_issuer:
                        cert_parts.append(cert_issuer)
                    if cert_date:
                        cert_parts.append(cert_date)
                    if cert_status and cert_status != 'Completed':
                        cert_parts.append(cert_status)

                    cert_text = " | ".join(cert_parts)
                    cert_para = doc.add_paragraph(cert_text, style='List Bullet')
                    cert_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    cert_para = doc.add_paragraph(str(cert), style='List Bullet')
                    cert_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ===== PROJECTS (NEW) =====
        if resume_data.get('projects'):
            doc.add_paragraph()
            doc.add_heading('PROJECTS', level=1)

            for project in resume_data['projects']:
                if isinstance(project, dict):
                    project_name = project.get('name', 'N/A')
                    technologies = project.get('technologies', [])
                    description = project.get('description', '')
                    github_url = project.get('github_url', '')

                    # Project name (bold)
                    project_para = doc.add_paragraph()
                    project_run = project_para.add_run(project_name)
                    project_run.bold = True
                    project_run.font.size = Pt(11)
                    project_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Technologies
                    if technologies:
                        tech_text = ", ".join(technologies) if isinstance(technologies, list) else str(technologies)
                        tech_para = doc.add_paragraph(f"Technologies: {tech_text}")
                        tech_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Description
                    if description:
                        desc_para = doc.add_paragraph(description, style='List Bullet')
                        desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # GitHub URL if available
                    if github_url:
                        github_para = doc.add_paragraph(github_url)
                        github_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Spacing between projects
                    doc.add_paragraph()

        # ===== LANGUAGES (NEW) =====
        if resume_data.get('languages'):
            doc.add_paragraph()
            doc.add_heading('LANGUAGES', level=1)

            language_items = []
            for lang in resume_data['languages']:
                if isinstance(lang, dict):
                    language = lang.get('language', 'N/A')
                    proficiency = lang.get('proficiency', '')
                    lang_text = f"{language}: {proficiency}" if proficiency else language
                    language_items.append(lang_text)
                else:
                    language_items.append(str(lang))

            # Display all languages on one line
            if language_items:
                lang_para = doc.add_paragraph(" | ".join(language_items))
                lang_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ===== DOCUMENT METADATA (for ATS) =====
        core_properties = doc.core_properties
        core_properties.title = f"Resume - {personal_info.get('full_name', 'N/A')}"
        core_properties.author = personal_info.get('full_name', 'N/A')
        core_properties.subject = "Professional Resume"

        # Add keywords from skills for ATS
        keywords = []
        if isinstance(resume_data.get('skills'), dict):
            for category, skills_data in resume_data['skills'].items():
                if isinstance(skills_data, dict):
                    keywords.extend(list(skills_data.keys())[:10])
                elif isinstance(skills_data, list):
                    keywords.extend(skills_data[:10])
        core_properties.keywords = ", ".join(keywords[:20])  # Max 20 keywords

        # Save document
        doc.save(output_path)

        print(f"✅ ATS-optimized DOCX generated: {output_path}")
        print(f"   - No tables, images, or complex formatting")
        print(f"   - Standard Calibri 11pt font")
        print(f"   - Simple bullets and left-aligned text")
        print(f"   - Metadata: {len(keywords)} keywords added")

        return output_path

    def _generate_classic_pdf(self, resume_data: Dict, output_path: str) -> str:
        """Generate classic PDF - similar to modern but with traditional styling"""
        # For now, use modern template
        # You can customize this later with more conservative styling
        return self._generate_modern_pdf(resume_data, output_path)

    def _generate_classic_docx(self, resume_data: Dict, output_path: str) -> str:
        """Generate classic DOCX"""
        return self._generate_modern_docx(resume_data, output_path)

    def _generate_technical_pdf(self, resume_data: Dict, output_path: str) -> str:
        """Generate technical PDF - emphasizes skills and projects"""
        return self._generate_modern_pdf(resume_data, output_path)

    def _generate_technical_docx(self, resume_data: Dict, output_path: str) -> str:
        """Generate technical DOCX"""
        return self._generate_modern_docx(resume_data, output_path)

    def _generate_from_docx_template(
        self,
        resume_data: Dict,
        output_path: str,
        template_key: str = "sales"
    ) -> str:
        """
        Generate resume using a pre-made DOCX template.

        This method loads a template file and fills it with the resume data.
        It uses a simple find-and-replace approach for placeholders.

        Args:
            resume_data: Structured resume data
            output_path: Path to save the generated DOCX
            template_key: Key for template (accounting, sales, attorney, office_manager)

        Returns:
            str: Path to generated DOCX
        """
        # Get template file path
        template_filename = self.docx_templates.get(template_key, self.docx_templates["sales"])
        template_path = os.path.join(self.templates_dir, template_filename)

        if not os.path.exists(template_path):
            print(f"Warning: Template not found at {template_path}")
            print("Falling back to custom generation")
            return self._generate_modern_docx(resume_data, output_path)

        print(f"Using template: {template_path}")

        try:
            # Load template
            doc = Document(template_path)

            # Extract resume data
            personal_info = resume_data.get("personal_info", {})
            professional_summary = resume_data.get("professional_summary", "")
            experiences = resume_data.get("work_experience", [])
            education = resume_data.get("education", [])
            skills = resume_data.get("skills", {})

            # Replace placeholders in the document
            # This is a simplified approach - you may need to customize based on actual template structure
            replacements = {
                "[Your Name]": personal_info.get("full_name", ""),
                "[Your Email]": personal_info.get("email", ""),
                "[Your Phone]": personal_info.get("phone", ""),
                "[Your Location]": personal_info.get("location", ""),
                "[LinkedIn URL]": personal_info.get("linkedin", ""),
                "[Professional Summary]": professional_summary,
            }

            # Replace text in paragraphs
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        # Replace inline
                        for run in paragraph.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, value)

            # Replace in tables (many templates use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in replacements.items():
                            if key in cell.text:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if key in run.text:
                                            run.text = run.text.replace(key, value)

            # Clear existing content sections and add new content
            # Note: This is a simplified implementation
            # You may need to customize based on the actual template structure

            # Save the modified document
            doc.save(output_path)
            print(f"Generated resume using template: {output_path}")
            return output_path

        except Exception as e:
            print(f"Error using template: {e}")
            print("Falling back to custom generation")
            return self._generate_modern_docx(resume_data, output_path)


# Convenience functions
def generate_resume_pdf(resume_data: Dict, output_path: str, template: str = "modern") -> str:
    """Generate PDF resume"""
    generator = ATSTemplateGenerator()
    return generator.generate_pdf(resume_data, output_path, template)


def generate_resume_docx(resume_data: Dict, output_path: str, template: str = "modern") -> str:
    """Generate DOCX resume"""
    generator = ATSTemplateGenerator()
    return generator.generate_docx(resume_data, output_path, template)
