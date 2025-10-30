"""
Resume Parser Service
Extracts text from uploaded PDF and DOCX resume files.
"""
import io
from typing import BinaryIO
from PyPDF2 import PdfReader
from docx import Document


class ResumeParserError(Exception):
    """Custom exception for resume parsing errors"""
    pass


class ResumeParser:
    """Service for parsing resume files (PDF/DOCX) into text"""

    @staticmethod
    def parse_pdf(file: BinaryIO) -> str:
        """
        Extract text from PDF file.

        Args:
            file: Binary file object (PDF)

        Returns:
            Extracted text content

        Raises:
            ResumeParserError: If PDF parsing fails
        """
        try:
            # Read PDF from binary stream
            pdf_reader = PdfReader(file)
            text_content = []

            # Extract text from all pages
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)

            if not text_content:
                raise ResumeParserError("No text content found in PDF")

            return "\n\n".join(text_content).strip()

        except Exception as e:
            raise ResumeParserError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_docx(file: BinaryIO) -> str:
        """
        Extract text from DOCX file.

        Args:
            file: Binary file object (DOCX)

        Returns:
            Extracted text content

        Raises:
            ResumeParserError: If DOCX parsing fails
        """
        try:
            # Read DOCX from binary stream
            document = Document(file)
            text_content = []

            # Extract text from all paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            if not text_content:
                raise ResumeParserError("No text content found in DOCX")

            return "\n".join(text_content).strip()

        except Exception as e:
            raise ResumeParserError(f"Failed to parse DOCX: {str(e)}")

    @classmethod
    def parse_resume(cls, file: BinaryIO, filename: str) -> str:
        """
        Parse resume file based on extension.

        Args:
            file: Binary file object
            filename: Original filename (used to determine type)

        Returns:
            Extracted text content

        Raises:
            ResumeParserError: If file type is unsupported or parsing fails
        """
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            return cls.parse_pdf(file)
        elif filename_lower.endswith('.docx'):
            return cls.parse_docx(file)
        else:
            raise ResumeParserError(
                f"Unsupported file type. Supported formats: PDF, DOCX. Got: {filename}"
            )
