"""
Advanced Template Handler for DOCX Resume Templates

This module provides intelligent template parsing and content replacement
while preserving the original template's typography, formatting, and style.
"""

import os
import re
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


class TemplateHandler:
    """
    Handles DOCX template parsing, placeholder detection, and smart content replacement.
    """

    # Common placeholder patterns to detect
    PLACEHOLDER_PATTERNS = [
        r'\[.*?\]',  # [Name], [Email], etc.
        r'\{.*?\}',  # {Name}, {Email}, etc.
        r'<<.*?>>',  # <<Name>>, <<Email>>, etc.
        r'YOUR\s+\w+',  # YOUR NAME, YOUR EMAIL, etc.
        r'Full\s+Name',
        r'Email\s+Address',
        r'Phone\s+Number',
        r'LinkedIn\s+URL',
        r'Professional\s+Summary',
        r'Work\s+Experience',
        r'Education',
        r'Skills',
        r'Company\s+Name',
        r'Job\s+Title',
        r'Start\s+Date',
        r'End\s+Date',
    ]

    # Mapping of placeholder types to resume data fields
    FIELD_MAPPINGS = {
        'name': ['name', 'full name', 'your name', 'full_name'],
        'email': ['email', 'email address', 'your email'],
        'phone': ['phone', 'phone number', 'telephone', 'your phone'],
        'linkedin': ['linkedin', 'linkedin url', 'linkedin profile'],
        'location': ['location', 'address', 'city', 'your location'],
        'summary': ['professional summary', 'summary', 'profile', 'about', 'objective'],
        'experience': ['work experience', 'experience', 'employment', 'work history'],
        'education': ['education', 'academic', 'qualifications'],
        'skills': ['skills', 'technical skills', 'core competencies'],
    }

    def __init__(self, templates_dir: Optional[str] = None):
        """
        Initialize template handler.

        Args:
            templates_dir: Directory containing DOCX templates
        """
        self.templates_dir = templates_dir or "/app/teamplate"

    def scan_templates(self) -> List[Dict[str, Any]]:
        """
        Scan templates directory and return available templates with metadata.

        Returns:
            List of template metadata dictionaries
        """
        templates = []

        if not os.path.exists(self.templates_dir):
            print(f"Warning: Templates directory not found: {self.templates_dir}")
            return templates

        for filename in os.listdir(self.templates_dir):
            if filename.endswith('.docx') and not filename.startswith('~'):
                template_path = os.path.join(self.templates_dir, filename)

                # Extract template info
                template_id = filename.replace('.docx', '').lower().replace(' ', '_')
                template_name = filename.replace('.docx', '')

                # Detect template type from name
                template_type = self._detect_template_type(filename)

                templates.append({
                    'id': template_id,
                    'name': template_name,
                    'filename': filename,
                    'path': template_path,
                    'type': template_type,
                })

        return templates

    def _detect_template_type(self, filename: str) -> str:
        """
        Detect template type from filename.

        Args:
            filename: Template filename

        Returns:
            Template type string
        """
        filename_lower = filename.lower()

        if 'accounting' in filename_lower:
            return 'accounting'
        elif 'sales' in filename_lower:
            return 'sales'
        elif 'attorney' in filename_lower or 'legal' in filename_lower:
            return 'legal'
        elif 'manager' in filename_lower:
            return 'management'
        elif 'technical' in filename_lower or 'engineer' in filename_lower:
            return 'technical'
        elif 'modern' in filename_lower:
            return 'modern'
        elif 'classic' in filename_lower:
            return 'classic'
        else:
            return 'general'

    def analyze_template(self, template_path: str) -> Dict[str, Any]:
        """
        Analyze a template to detect its structure and placeholders.

        Args:
            template_path: Path to the template file

        Returns:
            Dictionary with template analysis
        """
        doc = Document(template_path)

        analysis = {
            'placeholders': [],
            'sections': [],
            'styles_used': set(),
            'has_tables': False,
            'has_headers': False,
            'paragraph_count': 0,
        }

        # Analyze paragraphs
        for para in doc.paragraphs:
            analysis['paragraph_count'] += 1

            if para.style:
                analysis['styles_used'].add(para.style.name)

            # Detect placeholders
            placeholders = self._detect_placeholders(para.text)
            for placeholder in placeholders:
                analysis['placeholders'].append({
                    'text': placeholder,
                    'type': self._classify_placeholder(placeholder),
                    'location': 'paragraph'
                })

        # Analyze tables
        if doc.tables:
            analysis['has_tables'] = True
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            placeholders = self._detect_placeholders(para.text)
                            for placeholder in placeholders:
                                analysis['placeholders'].append({
                                    'text': placeholder,
                                    'type': self._classify_placeholder(placeholder),
                                    'location': 'table'
                                })

        # Analyze headers/footers
        for section in doc.sections:
            if section.header.paragraphs:
                analysis['has_headers'] = True

        analysis['styles_used'] = list(analysis['styles_used'])

        return analysis

    def _detect_placeholders(self, text: str) -> List[str]:
        """
        Detect placeholders in text using regex patterns.

        Args:
            text: Text to search

        Returns:
            List of detected placeholders
        """
        placeholders = []

        for pattern in self.PLACEHOLDER_PATTERNS:
            matches = re.findall(pattern, text, re.IGNORECASE)
            placeholders.extend(matches)

        return list(set(placeholders))  # Remove duplicates

    def _classify_placeholder(self, placeholder: str) -> str:
        """
        Classify a placeholder to determine what data it should be replaced with.

        Args:
            placeholder: Placeholder text

        Returns:
            Placeholder type (name, email, etc.)
        """
        placeholder_lower = placeholder.lower()

        for field_type, keywords in self.FIELD_MAPPINGS.items():
            for keyword in keywords:
                if keyword in placeholder_lower:
                    return field_type

        return 'unknown'

    def fill_template(
        self,
        template_path: str,
        resume_data: Dict,
        output_path: str
    ) -> str:
        """
        Fill a template with resume data while preserving formatting.

        Args:
            template_path: Path to template file
            resume_data: Resume data dictionary
            output_path: Path to save filled template

        Returns:
            Path to generated file
        """
        doc = Document(template_path)

        # Extract data from resume_data
        personal_info = resume_data.get('personal_info', {})
        work_experience = resume_data.get('work_experience', [])
        education = resume_data.get('education', [])
        skills = resume_data.get('skills', {})
        summary = resume_data.get('professional_summary', '')

        # Create replacement mappings
        replacements = self._build_replacements(resume_data)

        # Replace in paragraphs
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

        # Handle structured sections (experience, education, skills)
        self._fill_structured_sections(doc, resume_data)

        # Save filled template
        doc.save(output_path)

        print(f"✅ Template filled successfully: {output_path}")
        return output_path

    def _build_replacements(self, resume_data: Dict) -> Dict[str, str]:
        """
        Build a dictionary of placeholder -> value replacements.

        Args:
            resume_data: Resume data

        Returns:
            Dictionary mapping placeholders to values
        """
        personal_info = resume_data.get('personal_info', {})

        replacements = {}

        # Common placeholder variations
        name_placeholders = [
            '[Your Name]', '[Name]', '[Full Name]', 'YOUR NAME', 'Full Name',
            '{Name}', '{YOUR NAME}', '<<Name>>', '[FULL NAME]'
        ]
        for placeholder in name_placeholders:
            replacements[placeholder] = personal_info.get('full_name', '')

        email_placeholders = [
            '[Email]', '[Email Address]', '[Your Email]', 'YOUR EMAIL', 'Email Address',
            '{Email}', '<<Email>>', '[EMAIL]'
        ]
        for placeholder in email_placeholders:
            replacements[placeholder] = personal_info.get('email', '')

        phone_placeholders = [
            '[Phone]', '[Phone Number]', '[Your Phone]', 'YOUR PHONE', 'Phone Number',
            '{Phone}', '<<Phone>>', '[PHONE NUMBER]'
        ]
        for placeholder in phone_placeholders:
            replacements[placeholder] = personal_info.get('phone', '')

        linkedin_placeholders = [
            '[LinkedIn]', '[LinkedIn URL]', '[LinkedIn Profile]', 'YOUR LINKEDIN',
            '{LinkedIn}', '<<LinkedIn>>'
        ]
        for placeholder in linkedin_placeholders:
            replacements[placeholder] = personal_info.get('linkedin', '')

        location_placeholders = [
            '[Location]', '[Address]', '[City]', '[Your Location]', 'YOUR LOCATION',
            '{Location}', '<<Location>>'
        ]
        for placeholder in location_placeholders:
            replacements[placeholder] = personal_info.get('location', '')

        summary_placeholders = [
            '[Professional Summary]', '[Summary]', '[Profile]', 'YOUR SUMMARY',
            '{Summary}', '<<Summary>>'
        ]
        for placeholder in summary_placeholders:
            replacements[placeholder] = resume_data.get('professional_summary', '')

        return replacements

    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        """
        Replace placeholders in a paragraph while preserving formatting.

        Args:
            paragraph: Paragraph object
            replacements: Dictionary of replacements
        """
        # Replace in full text first (simple case)
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                # For simple replacements, iterate through runs
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)

    def _fill_structured_sections(self, doc: Document, resume_data: Dict):
        """
        Fill structured sections like work experience and education.

        This is a simplified version - you may need to customize based on
        your specific templates.

        Args:
            doc: Document object
            resume_data: Resume data
        """
        # This is where you'd implement logic to find and fill
        # repeating sections like work experience entries
        # For now, this is a placeholder for future enhancement
        pass


def get_available_templates() -> List[Dict[str, Any]]:
    """
    Convenience function to get available templates.

    Returns:
        List of template metadata
    """
    handler = TemplateHandler()
    return handler.scan_templates()


def analyze_template_file(template_path: str) -> Dict[str, Any]:
    """
    Convenience function to analyze a template.

    Args:
        template_path: Path to template file

    Returns:
        Template analysis
    """
    handler = TemplateHandler()
    return handler.analyze_template(template_path)


if __name__ == "__main__":
    """
    Test the template handler
    """
    print("="*80)
    print("TEMPLATE HANDLER TEST")
    print("="*80)

    handler = TemplateHandler()

    # Scan templates
    print("\n📁 Scanning templates...")
    templates = handler.scan_templates()

    print(f"\n✅ Found {len(templates)} templates:")
    for template in templates:
        print(f"\n  Template: {template['name']}")
        print(f"    ID: {template['id']}")
        print(f"    Type: {template['type']}")
        print(f"    Path: {template['path']}")

        # Analyze template
        print(f"\n  📊 Analyzing template...")
        analysis = handler.analyze_template(template['path'])
        print(f"    Paragraphs: {analysis['paragraph_count']}")
        print(f"    Has tables: {analysis['has_tables']}")
        print(f"    Has headers: {analysis['has_headers']}")
        print(f"    Styles used: {', '.join(analysis['styles_used'][:5])}")
        print(f"    Placeholders detected: {len(analysis['placeholders'])}")

        if analysis['placeholders']:
            print(f"    Sample placeholders:")
            for ph in analysis['placeholders'][:5]:
                print(f"      - {ph['text']} (type: {ph['type']}, location: {ph['location']})")
